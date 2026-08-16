import asyncio
import json
import os
import zipfile
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import httpx
import pytest
from fastapi.testclient import TestClient

from app.config import Settings
from app import agent as agent_module
from app.auth import AuthIdentity, get_auth_identity
from app.database import SessionLocal
from app.main import app, invitation_token_hash, settings as api_settings
from app.models import AuditEvent, Invitation, Membership, OrganizationModelPolicy, UserSettings
from app.knowledge_worker import run_once as run_ingestion_once
from app.artifact_worker import run_once as run_artifact_once
from app.observability import RequestLoggingMiddleware, bind_request_context, configure_logging, exception_stack, log_transcript
from sqlalchemy import select


OWNER = {"X-User-ID": "user-ayman", "X-Organization-ID": "org-northstar"}
ADMIN = {"X-User-ID": "user-maya", "X-Organization-ID": "org-northstar"}
MEMBER = {"X-User-ID": "user-jon", "X-Organization-ID": "org-northstar"}


def test_health_and_seeded_identity():
    with TestClient(app) as client:
        health = client.get("/health", headers={"X-Request-ID": "health-check-123"})
        assert health.json()["status"] == "ok"
        assert health.headers["X-Request-ID"] == "health-check-123"
        response = client.get("/v1/me", headers=OWNER)
        assert response.status_code == 200
        assert response.json()["role"] == "owner"
        models = client.get("/v1/models", headers=OWNER)
        assert [(item["display_name"], item["id"]) for item in models.json()["models"]] == [
            ("Default", "gemini-3.5-flash"),
            ("Pro", "gemini-3.1-pro-preview"),
        ]


def test_invalid_request_id_is_replaced():
    with TestClient(app) as client:
        response = client.get("/health", headers={"X-Request-ID": "bad request id"})
    assert response.status_code == 200
    assert response.headers["X-Request-ID"] != "bad request id"
    assert len(response.headers["X-Request-ID"]) == 36


@pytest.mark.asyncio
async def test_concurrent_request_context_is_isolated():
    async def echo_context(scope, receive, send):
        del receive
        bind_request_context(user_id=scope["path"].removeprefix("/"))
        await asyncio.sleep(0.01)
        await send({"type": "http.response.start", "status": 204, "headers": []})
        await send({"type": "http.response.body", "body": b""})

    transport = httpx.ASGITransport(app=RequestLoggingMiddleware(echo_context))
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        first, second = await asyncio.gather(
            client.get("/first", headers={"X-Request-ID": "request-first"}),
            client.get("/second", headers={"X-Request-ID": "request-second"}),
        )
    assert first.headers["X-Request-ID"] == "request-first"
    assert second.headers["X-Request-ID"] == "request-second"


def test_private_conversation_isolation():
    with TestClient(app) as client:
        assert client.get("/v1/conversations/conversation-quarterly", headers=MEMBER).status_code == 404


def test_member_cannot_mutate_shared_prompts_but_can_read():
    with TestClient(app) as client:
        assert client.get("/v1/prompts", headers=MEMBER).status_code == 200
        response = client.post("/v1/prompts", headers=MEMBER, json={"title": "Unauthorized", "body": "No", "tags": []})
        assert response.status_code == 403


def test_admin_creates_versioned_prompt():
    with TestClient(app) as client:
        created = client.post("/v1/prompts", headers=ADMIN, json={"title": "Decision memo", "description": "Create a decision record", "body": "Summarize the decision and rationale.", "tags": ["Leadership"]})
        assert created.status_code == 201
        prompt = created.json()
        updated = client.patch(f"/v1/prompts/{prompt['id']}", headers=ADMIN, json={"body": "Summarize the decision, rationale, owner, and review date."})
        assert updated.status_code == 200
        assert updated.json()["version_number"] == 2


def test_invitation_can_be_accepted_before_membership_exists():
    with TestClient(app) as client:
        invited = client.post("/v1/organizations/current/invitations", headers=OWNER, json={"email": "priya@northstaradvisory.com"})
        assert invited.status_code == 201
        token = invited.json()["acceptance_token"]
        preview = client.get(f"/v1/invitations/{token}/preview")
        assert preview.status_code == 200
        assert preview.json()["organization"]["name"] == "Northstar Advisory"
        assert preview.json()["invited_email"] == "p••••@northstaradvisory.com"
        accepted = client.post(f"/v1/invitations/{token}/accept", headers={"X-User-ID": "user-priya"})
        assert accepted.status_code == 200
        assert accepted.json()["organization_id"] == "org-northstar"
        assert accepted.json()["organization"]["role"] == "member"
        assert client.post(f"/v1/invitations/{token}/accept", headers={"X-User-ID": "user-priya"}).status_code == 200


def test_invitation_tokens_are_hashed_and_rotated():
    async def invitation_row(invitation_id: str):
        async with SessionLocal() as db:
            return await db.get(Invitation, invitation_id)

    with TestClient(app) as client:
        first = client.post("/v1/organizations/current/invitations", headers=OWNER, json={"email": "rotated@example.com"})
        assert first.status_code == 201
        first_body = first.json()
        first_token = first_body["acceptance_token"]
        row = asyncio.run(invitation_row(first_body["id"]))
        assert row.token_hash == invitation_token_hash(first_token)
        assert not hasattr(row, "token")

        second = client.post("/v1/organizations/current/invitations", headers=OWNER, json={"email": "ROTATED@example.com"})
        assert second.status_code == 201
        second_body = second.json()
        assert second_body["id"] == first_body["id"]
        assert second_body["acceptance_token"] != first_token
        assert client.get(f"/v1/invitations/{first_token}/preview").status_code == 404
        assert client.get(f"/v1/invitations/{second_body['acceptance_token']}/preview").status_code == 200
        api_log = (Path(os.environ["LOG_DIR"]) / "api.jsonl").read_text()
        assert first_token not in api_log
        assert second_body["acceptance_token"] not in api_log
        assert "/v1/invitations/[token]/preview" in api_log


def test_firebase_bootstrap_verification_and_atomic_organization_creation():
    identity = AuthIdentity(uid="firebase-new-user", email="new.owner@example.com", email_verified=False, display_name="New Owner")

    async def override_identity():
        return identity

    previous_mode = api_settings.auth_mode
    api_settings.auth_mode = "firebase"
    app.dependency_overrides[get_auth_identity] = override_identity
    try:
        with TestClient(app) as client:
            bootstrapped = client.post("/v1/auth/bootstrap", json={"display_name": "New Owner"})
            assert bootstrapped.status_code == 200
            assert bootstrapped.json()["requires_onboarding"] is True
            assert bootstrapped.json()["email_verified"] is False
            assert client.post("/v1/organizations", json={"name": "Unverified Workspace"}).status_code == 403

            identity = AuthIdentity(uid="firebase-new-user", email="new.owner@example.com", email_verified=True, display_name="New Owner")
            created = client.post("/v1/organizations", json={"name": "Verified Workspace"})
            assert created.status_code == 201
            organization_id = created.json()["id"]
            refreshed = client.post("/v1/auth/bootstrap")
            assert refreshed.json()["requires_onboarding"] is False
            assert refreshed.json()["organizations"][0]["role"] == "owner"

        async def assert_workspace_defaults():
            async with SessionLocal() as db:
                user_id = (await db.scalar(select(Membership.user_id).where(Membership.organization_id == organization_id)))
                assert await db.scalar(select(UserSettings.id).where(UserSettings.organization_id == organization_id, UserSettings.user_id == user_id))
                assert await db.scalar(select(OrganizationModelPolicy.id).where(OrganizationModelPolicy.organization_id == organization_id))
                assert await db.scalar(select(AuditEvent.id).where(AuditEvent.organization_id == organization_id, AuditEvent.action == "organization.created"))

        asyncio.run(assert_workspace_defaults())
    finally:
        app.dependency_overrides.pop(get_auth_identity, None)
        api_settings.auth_mode = previous_mode


def test_demo_stream_emits_typed_events():
    with TestClient(app) as client:
        with client.stream("POST", "/v1/conversations/conversation-quarterly/messages/stream", headers=OWNER, json={"content": "Give me the next action", "model": "gemini-3.1-pro-preview", "effort": "medium", "attachment_ids": []}) as response:
            body = "".join(response.iter_text())
        assert response.status_code == 200
        assert "event: message_started" in body
        assert "event: text_delta" in body
        assert "event: message_completed" in body


def test_development_transcript_is_separate_from_general_log():
    secret = "PRIVATE_CHAT_SENTINEL_7e72d4"
    with TestClient(app) as client:
        with client.stream(
            "POST",
            "/v1/conversations/conversation-quarterly/messages/stream",
            headers={**OWNER, "X-Request-ID": "transcript-separation-test"},
            json={"content": secret, "model": "gemini-3.1-pro-preview", "effort": "medium", "attachment_ids": []},
        ) as response:
            assert "event: message_completed" in "".join(response.iter_text())

    log_dir = Path(os.environ["LOG_DIR"])
    api_log = (log_dir / "api.jsonl").read_text()
    transcript_log = (log_dir / "chat-transcripts.jsonl").read_text()
    assert secret not in api_log
    assert secret in transcript_log
    record = next(json.loads(line) for line in transcript_log.splitlines() if secret in line)
    assert record["outcome"] == "completed"
    assert record["request_id"] == "transcript-separation-test"


def test_attachment_logs_metadata_without_file_contents():
    file_secret = b"PRIVATE_FILE_BYTES_SENTINEL_291f"
    with TestClient(app) as client:
        response = client.post(
            "/v1/conversations/conversation-quarterly/attachments",
            headers={**OWNER, "X-Request-ID": "attachment-logging-test"},
            files={"upload": ("debug-notes.txt", file_secret, "text/plain")},
        )
    assert response.status_code == 201
    api_log = (Path(os.environ["LOG_DIR"]) / "api.jsonl").read_text()
    assert "attachment.upload_completed" in api_log
    assert "debug-notes.txt" in api_log
    assert file_secret.decode() not in api_log


def test_transcripts_are_forcibly_disabled_outside_development(tmp_path):
    production_log_dir = tmp_path / "production-logs"
    try:
        configure_logging(Settings(app_env="production", log_dir=production_log_dir, log_chat_transcripts=True))
        log_transcript(user_message="must-not-be-written", assistant_response="also-private")
        assert not (production_log_dir / "chat-transcripts.jsonl").exists()
    finally:
        configure_logging(Settings(log_dir=Path(os.environ["LOG_DIR"])))


def test_sanitized_exception_stack_omits_exception_message():
    try:
        raise RuntimeError("PRIVATE_EXCEPTION_SENTINEL")
    except RuntimeError as exc:
        rendered = "\n".join(exception_stack(exc))
    assert "PRIVATE_EXCEPTION_SENTINEL" not in rendered
    assert "test_sanitized_exception_stack_omits_exception_message" in rendered


def test_google_provider_passes_dotenv_key_to_client(monkeypatch):
    api_key = "test-api-key-not-a-real-secret"
    monkeypatch.setattr(
        agent_module,
        "get_settings",
        lambda: SimpleNamespace(gemini_model="gemini-3.5-flash", google_api_key=api_key),
    )
    provider = agent_module.GoogleAdkChatProvider()
    assert provider.agent.model.model == "gemini-3.5-flash"
    assert provider.agent.model.client_kwargs == {"api_key": api_key}


def test_owner_must_transfer_before_account_deletion():
    with TestClient(app) as client:
        response = client.delete("/v1/me", headers={"X-User-ID": "user-ayman"})
        assert response.status_code == 409


def test_non_owner_can_delete_personal_account():
    with TestClient(app) as client:
        response = client.delete("/v1/me", headers={"X-User-ID": "user-priya"})
        assert response.status_code == 204
        assert client.get("/v1/organizations", headers={"X-User-ID": "user-priya"}).status_code == 401


def test_knowledge_permissions_ingestion_retrieval_and_scope_snapshots():
    with TestClient(app) as client:
        created = client.post("/v1/knowledge-bases", headers=OWNER, json={
            "title": "Employee handbook test",
            "description": "Private HR practices",
            "member_ids": ["user-jon"],
        })
        assert created.status_code == 201
        knowledge_base_id = created.json()["id"]

        renamed = client.patch(
            f"/v1/knowledge-bases/{knowledge_base_id}",
            headers=OWNER,
            json={"title": "People handbook test", "description": "Current people policies and practices"},
        )
        assert renamed.status_code == 200
        assert renamed.json()["title"] == "People handbook test"
        assert renamed.json()["description"] == "Current people policies and practices"
        assert client.patch(
            f"/v1/knowledge-bases/{knowledge_base_id}",
            headers=MEMBER,
            json={"title": "Unauthorized rename"},
        ).status_code == 403

        assert any(item["id"] == knowledge_base_id for item in client.get("/v1/knowledge-bases", headers=MEMBER).json())
        assert client.get(f"/v1/knowledge-bases/{knowledge_base_id}", headers=ADMIN).status_code == 404
        managed = client.get("/v1/knowledge-bases-management", headers=ADMIN)
        assert any(item["id"] == knowledge_base_id and item["has_access"] is False for item in managed.json())
        self_grant = client.post(f"/v1/knowledge-bases/{knowledge_base_id}/self-grant", headers=ADMIN, json={"reason": "Reviewing ingestion quality"})
        assert self_grant.status_code == 201

        uploaded = client.post(
            f"/v1/knowledge-bases/{knowledge_base_id}/documents",
            headers=MEMBER,
            files=[("uploads", ("leave-policy.md", b"Employees receive 20 paid vacation days each calendar year. Carryover requires manager approval.", "text/markdown"))],
        )
        assert uploaded.status_code == 202
        asyncio.run(run_ingestion_once())

        detail = client.get(f"/v1/knowledge-bases/{knowledge_base_id}", headers=MEMBER).json()
        assert detail["documents"][0]["versions"][0]["extraction_status"] == "ready"
        search = client.get("/v1/knowledge/search", headers=MEMBER, params=[("q", "How many vacation days?"), ("knowledge_base_ids", knowledge_base_id)])
        assert search.status_code == 200
        assert search.json()["results"][0]["title"] == "leave-policy.md"

        exact_duplicate = client.post(
            f"/v1/knowledge-bases/{knowledge_base_id}/documents",
            headers=MEMBER,
            files=[("uploads", ("copy.md", b"Employees receive 20 paid vacation days each calendar year. Carryover requires manager approval.", "text/markdown"))],
        )
        assert exact_duplicate.status_code == 409

        conversation = client.post("/v1/conversations", headers=MEMBER, json={
            "title": "Policy question",
            "knowledge_base_ids": [knowledge_base_id],
            "web_search_enabled": False,
        })
        assert conversation.status_code == 201
        conversation_id = conversation.json()["id"]
        with client.stream("POST", f"/v1/conversations/{conversation_id}/messages/stream", headers=MEMBER, json={
            "content": "How many vacation days do employees receive?",
            "knowledge_base_ids": [knowledge_base_id],
            "web_search_enabled": False,
        }) as response:
            body = "".join(response.iter_text())
        assert "event: retrieval_started" in body
        assert "event: internal_citations" in body
        loaded = client.get(f"/v1/conversations/{conversation_id}", headers=MEMBER).json()
        assert loaded["messages"][-1]["knowledge_base_ids"] == [knowledge_base_id]
        assert loaded["messages"][-1]["citations"][0]["source_type"] == "company"


def test_web_default_and_ambiguous_company_question_emit_explainable_state():
    with TestClient(app) as client:
        assert client.get("/v1/settings", headers=OWNER).json()["web_search_default"] is True
        updated = client.patch("/v1/settings", headers=OWNER, json={"web_search_default": True})
        assert updated.json()["web_search_default"] is True
        conversation = client.post("/v1/conversations", headers=OWNER, json={"title": "Research"})
        assert conversation.status_code == 201
        assert conversation.json()["web_search_enabled"] is True

        first = client.post("/v1/knowledge-bases", headers=OWNER, json={"title": "Ambiguity one", "member_ids": []}).json()
        second = client.post("/v1/knowledge-bases", headers=OWNER, json={"title": "Ambiguity two", "member_ids": []}).json()
        scoped = client.post("/v1/conversations", headers=OWNER, json={"title": "Ambiguous", "knowledge_base_ids": [first["id"], second["id"]], "web_search_enabled": False}).json()
        with client.stream("POST", f"/v1/conversations/{scoped['id']}/messages/stream", headers=OWNER, json={
            "content": "What is the policy?",
            "knowledge_base_ids": [first["id"], second["id"]],
            "web_search_enabled": False,
        }) as response:
            body = "".join(response.iter_text())
        assert "event: clarification_required" in body
        assert "Which team, region, time period" in body


def test_deep_research_mode_is_available_without_an_artifact_and_requires_web_search():
    with TestClient(app) as client:
        conversation = client.post("/v1/conversations", headers=OWNER, json={
            "title": "Deep research chat",
            "knowledge_base_ids": [],
            "web_search_enabled": False,
        }).json()
        with client.stream("POST", f"/v1/conversations/{conversation['id']}/messages/stream", headers=OWNER, json={
            "content": "Evaluate this market opportunity",
            "knowledge_base_ids": [],
            "web_search_enabled": False,
            "research_mode": "deep",
        }) as response:
            body = "".join(response.iter_text())

        assert response.status_code == 200
        assert "event: message_completed" in body
        assert client.get(f"/v1/conversations/{conversation['id']}", headers=OWNER).json()["web_search_enabled"] is True


def _sse_payload(body: str, event_name: str) -> dict:
    for block in body.split("\n\n"):
        if f"event: {event_name}" not in block:
            continue
        data = next((line.removeprefix("data: ") for line in block.splitlines() if line.startswith("data: ")), None)
        if data:
            return json.loads(data)
    raise AssertionError(f"Missing SSE event: {event_name}")


@pytest.mark.parametrize(("format_name", "expected_entry"), [("docx", "word/document.xml"), ("pptx", "ppt/presentation.xml")])
def test_editable_artifact_generation_download_and_revision(format_name, expected_entry):
    with TestClient(app) as client:
        with client.stream("POST", "/v1/conversations/conversation-quarterly/messages/stream", headers=OWNER, json={
            "content": "Create a practical four-week study plan for a manager learning financial analysis.",
            "model": "gemini-3.5-flash",
            "effort": "medium",
            "attachment_ids": [],
            "knowledge_base_ids": [],
            "web_search_enabled": False,
            "artifact_request": {"format": format_name, "template_id": "auto", "use_document_template": True},
        }) as response:
            body = "".join(response.iter_text())
        assert response.status_code == 200
        queued = _sse_payload(body, "artifact_queued")["artifact"]
        assert queued["format"] == format_name
        assert queued["status"] == "queued"

        assert asyncio.run(run_artifact_once()) is True
        ready = client.get(f"/v1/artifacts/{queued['id']}", headers=OWNER)
        assert ready.status_code == 200
        assert ready.json()["status"] == "ready"
        assert ready.json()["version"]["qa"]["structural"] == "passed"
        download = client.get(f"/v1/artifacts/{queued['id']}/download", headers=OWNER)
        assert download.status_code == 200
        with zipfile.ZipFile(BytesIO(download.content)) as archive:
            assert expected_entry in archive.namelist()

        revised = client.post(f"/v1/artifacts/{queued['id']}/revisions", headers=OWNER, json={"instructions": "Add a weekly review checklist."})
        assert revised.status_code == 202
        assert revised.json()["current_version"] == 2
        assert asyncio.run(run_artifact_once()) is True
        versions = client.get(f"/v1/artifacts/{queued['id']}", headers=OWNER).json()["versions"]
        assert [item["version_number"] for item in versions] == [2, 1]
        assert all(item["status"] == "ready" for item in versions)


def test_artifact_natural_request_pdf_handling_template_permissions_and_cancellation():
    with TestClient(app) as client:
        assert client.post(
            "/v1/organizations/current/document-template",
            headers=MEMBER,
            files={"upload": ("template.docx", b"not-a-document", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        ).status_code == 403

        with client.stream("POST", "/v1/conversations/conversation-quarterly/messages/stream", headers=OWNER, json={"content": "Make a PDF study plan for calculus", "knowledge_base_ids": [], "web_search_enabled": False}) as response:
            pdf_body = "".join(response.iter_text())
        assert "PDF export is not available yet" in pdf_body
        assert "event: artifact_queued" not in pdf_body

        with client.stream("POST", "/v1/conversations/conversation-quarterly/messages/stream", headers=OWNER, json={"content": "Create a PowerPoint about customer onboarding", "knowledge_base_ids": [], "web_search_enabled": False}) as response:
            queued_body = "".join(response.iter_text())
        artifact = _sse_payload(queued_body, "artifact_queued")["artifact"]
        assert artifact["format"] == "pptx"
        cancelled = client.post(f"/v1/artifacts/{artifact['id']}/cancel", headers=OWNER)
        assert cancelled.status_code == 200
        assert cancelled.json()["status"] == "cancelled"
        assert client.get(f"/v1/artifacts/{artifact['id']}", headers=MEMBER).status_code == 404


def _word_template_bytes(*, multiple_sections: bool = False) -> bytes:
    from docx import Document
    from docx.enum.section import WD_SECTION

    document = Document()
    document.sections[0].header.paragraphs[0].text = "Northstar Advisory letterhead"
    document.sections[0].footer.paragraphs[0].text = "Confidential"
    document.add_paragraph("THIS SAMPLE BODY MUST BE REMOVED")
    if multiple_sections:
        document.add_section(WD_SECTION.NEW_PAGE)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_document_template_validation_activation_and_artifact_fidelity():
    with TestClient(app) as client:
        dotx = client.post(
            "/v1/organizations/current/document-template",
            headers=OWNER,
            files={"upload": ("letterhead.dotx", b"not-used", "application/vnd.openxmlformats-officedocument.wordprocessingml.template")},
        )
        assert dotx.status_code == 415
        assert "Word Document (.docx)" in dotx.json()["detail"]

        multiple = client.post(
            "/v1/organizations/current/document-template",
            headers=OWNER,
            files={"upload": ("multi.docx", _word_template_bytes(multiple_sections=True), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert multiple.status_code == 422
        assert "exactly one Word section" in multiple.json()["detail"]

        uploaded = client.post(
            "/v1/organizations/current/document-template",
            headers=OWNER,
            files={"upload": ("northstar-letterhead.docx", _word_template_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert uploaded.status_code == 202
        pending = uploaded.json()["pending_version"]
        assert pending["status"] == "queued"
        assert uploaded.json()["active_version"] is None

        assert asyncio.run(run_artifact_once()) is True
        template = client.get("/v1/organizations/current/document-template", headers=OWNER).json()
        assert template["enabled"] is True
        assert template["active_version"]["id"] == pending["id"]
        assert template["active_version"]["validation_report"]["sample_body_discarded"] is True
        assert client.get("/v1/organizations/current/document-template", headers=MEMBER).json()["can_manage"] is False

        with client.stream("POST", "/v1/conversations/conversation-quarterly/messages/stream", headers=OWNER, json={
            "content": "Create an editable document explaining a practical AI rollout.",
            "knowledge_base_ids": [],
            "web_search_enabled": False,
            "artifact_request": {"format": "docx", "template_id": "auto", "use_document_template": True},
        }) as response:
            queued = _sse_payload("".join(response.iter_text()), "artifact_queued")["artifact"]
        assert queued["version"]["document_template_version_id"] == pending["id"]
        assert asyncio.run(run_artifact_once()) is True
        ready = client.get(f"/v1/artifacts/{queued['id']}", headers=OWNER).json()
        assert ready["version"]["qa"]["manual_page_breaks"] == 0
        download = client.get(f"/v1/artifacts/{queued['id']}/download", headers=OWNER)
        with zipfile.ZipFile(BytesIO(download.content)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
            header_xml = "\n".join(archive.read(name).decode("utf-8") for name in archive.namelist() if name.startswith("word/header"))
        assert "THIS SAMPLE BODY MUST BE REMOVED" not in document_xml
        assert "Northstar Advisory letterhead" in header_xml
        assert 'w:type="page"' not in document_xml

        rejected_replacement = client.post(
            "/v1/organizations/current/document-template",
            headers=OWNER,
            files={"upload": ("invalid-replacement.docx", _word_template_bytes(multiple_sections=True), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert rejected_replacement.status_code == 422
        assert client.get("/v1/organizations/current/document-template", headers=OWNER).json()["active_version_id"] == pending["id"]

        with client.stream("POST", "/v1/conversations/conversation-quarterly/messages/stream", headers=OWNER, json={
            "content": "Create an editable unbranded document.",
            "knowledge_base_ids": [],
            "web_search_enabled": False,
            "artifact_request": {"format": "docx", "template_id": "auto", "use_document_template": False},
        }) as response:
            untemplated = _sse_payload("".join(response.iter_text()), "artifact_queued")["artifact"]
        assert untemplated["version"]["document_template_version_id"] is None
        assert client.post(f"/v1/artifacts/{untemplated['id']}/cancel", headers=OWNER).status_code == 200

        revised = client.post(f"/v1/artifacts/{queued['id']}/revisions", headers=OWNER, json={"instructions": "Shorten the introduction."})
        assert revised.status_code == 202
        assert revised.json()["version"]["document_template_version_id"] == pending["id"]

        disabled = client.post("/v1/organizations/current/document-template/disable", headers=ADMIN)
        assert disabled.status_code == 200
        assert disabled.json()["enabled"] is False
